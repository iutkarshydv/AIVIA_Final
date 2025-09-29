# Create the text extraction service (app/services/text_extractor.py)
text_extractor_py_content = '''"""
AIVIA MVP Simple Text Extraction Service
Extract text from PDF and DOCX resume files for knowledge base
"""
import pdfplumber
import docx
from typing import Optional, Union
import logging
import os
from pathlib import Path


class TextExtractionError(Exception):
    """Custom exception for text extraction failures."""
    pass


class SimpleTextExtractor:
    """
    Simple text extraction from resume files for knowledge base.
    Supports PDF and DOCX files with proper error handling.
    """
    
    def __init__(self):
        self.logger = logging.getLogger(__name__)
        self.supported_extensions = {'.pdf', '.docx', '.doc'}
        self.max_file_size = 10 * 1024 * 1024  # 10MB limit
    
    async def extract_text_from_file(self, file_path: str, file_type: str) -> str:
        """
        Extract raw text from resume file with validation and error handling.
        
        Args:
            file_path: Path to the resume file
            file_type: File type (pdf, docx, doc)
            
        Returns:
            str: Extracted text content
            
        Raises:
            TextExtractionError: If extraction fails or file is invalid
        """
        # Validation
        if not file_path or not os.path.exists(file_path):
            raise TextExtractionError(f"File not found: {file_path}")
        
        # Check file size
        file_size = os.path.getsize(file_path)
        if file_size > self.max_file_size:
            raise TextExtractionError(f"File too large: {file_size} bytes (max: {self.max_file_size})")
        
        if file_size == 0:
            raise TextExtractionError("File is empty")
        
        # Normalize file type
        file_type = file_type.lower().replace('.', '')
        
        try:
            if file_type == 'pdf':
                text = self._extract_from_pdf(file_path)
            elif file_type in ['docx', 'doc']:
                text = self._extract_from_docx(file_path)
            else:
                raise TextExtractionError(f"Unsupported file type: {file_type}")
            
            # Validate extracted text
            if not text or not text.strip():
                raise TextExtractionError("No text content found in file")
            
            # Clean and normalize text
            cleaned_text = self._clean_text(text)
            
            self.logger.info(f"Successfully extracted {len(cleaned_text)} characters from {file_path}")
            return cleaned_text
            
        except Exception as e:
            if isinstance(e, TextExtractionError):
                raise e
            raise TextExtractionError(f"Text extraction failed: {str(e)}")
    
    def _extract_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file with error handling."""
        text_parts = []
        
        try:
            with pdfplumber.open(file_path) as pdf:
                if not pdf.pages:
                    raise TextExtractionError("PDF contains no pages")
                
                for page_num, page in enumerate(pdf.pages, 1):
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
                    else:
                        self.logger.warning(f"No text found on page {page_num}")
                
        except Exception as e:
            raise TextExtractionError(f"PDF extraction failed: {str(e)}")
        
        return "\\n".join(text_parts)
    
    def _extract_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file with error handling."""
        text_parts = []
        
        try:
            doc = docx.Document(file_path)
            
            if not doc.paragraphs:
                raise TextExtractionError("DOCX contains no paragraphs")
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
                    
        except Exception as e:
            raise TextExtractionError(f"DOCX extraction failed: {str(e)}")
        
        return "\\n".join(text_parts)
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize extracted text."""
        if not text:
            return ""
        
        # Remove excessive whitespace
        lines = [line.strip() for line in text.split('\\n')]
        lines = [line for line in lines if line]  # Remove empty lines
        
        # Join with single newlines
        cleaned = '\\n'.join(lines)
        
        # Limit text length for knowledge base (ElevenLabs may have limits)
        max_length = 10000  # Conservative limit
        if len(cleaned) > max_length:
            cleaned = cleaned[:max_length] + "..."
            self.logger.warning(f"Text truncated to {max_length} characters")
        
        return cleaned
'''

with open('text_extractor.py', 'w') as f:
    f.write(text_extractor_py_content)

print("✅ Created text_extractor.py - Text extraction service")